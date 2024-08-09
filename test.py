import os
import requests
import pandas as pd
from bs4 import BeautifulSoup
from PIL import Image
from io import BytesIO
from docx import Document
from docx.shared import Inches

# 设置目标网址
url = 'https://imagingmall.com/jumpshop/EventPhoto?EventId=00025&page=2&SearchTab=0&IsResearch=True&FaceSortFlag=False#pagetop'  # 替换成你要抓取的实际网站URL

# 创建一个目录来保存图片
if not os.path.exists('images'):
    os.makedirs('images')

def download_image(img_url, img_number):
    try:
        # 下载图片
        response = requests.get(img_url)
        response.raise_for_status()
        
        # 读取图片
        img = Image.open(BytesIO(response.content))
        
        # 设置保存路径
        img_filename = os.path.join('images', f'image_{img_number}.jpg')
        
        # 保存图片
        img.save(img_filename)
        print(f"图片 {img_number} 保存成功: {img_filename}")
        return img_filename
    except Exception as e:
        print(f"下载图片失败: {e}")
        return None

def create_word_document(image_data, filename):
    doc = Document()
    doc.add_heading('Downloaded Images', level=1)
    
    for img_number, img_file in image_data:
        # 添加图片标题
        doc.add_heading(f'Image {img_number}', level=2)
        
        # 添加图片
        doc.add_picture(img_file, width=Inches(3))  # 设置图片宽度为 4 英寸
        
        # 添加图片的 URL
        # doc.add_paragraph(f'Image URL: {img_url}')
    
    doc.save(filename)
    print(f"Word 文件已保存: {filename}")

def main():
    # 获取网页内容
    response = requests.get(url)
    response.raise_for_status()
    
    # 解析网页
    soup = BeautifulSoup(response.text, 'html.parser')
    
    # 查找所有图片标签
    img_tags = soup.find_all('img')
    
    image_data = []
    img_number = 1
    for img_tag in img_tags:
        # 获取图片 URL
        img_url = img_tag.get('src')
        if not img_url:
            continue
        
        # 处理相对 URL
        if not img_url.startswith('http'):
            img_url = requests.compat.urljoin(url, img_url)
        
        # 下载图片并保存
        img_file = download_image(img_url, img_number)
        if img_file:
            image_data.append((img_number, img_file))
        img_number += 1

    
    # 创建 Word 文件
    create_word_document(image_data, 'image_data_2.docx')

if __name__ == '__main__':
    main()
